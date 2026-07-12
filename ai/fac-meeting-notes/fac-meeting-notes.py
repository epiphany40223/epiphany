#!/usr/bin/env python3
"""Generate Finance Advisory Committee meeting minutes from source documents.

Reads a meeting transcript plus the Financial Highlights, Financial
Observations, and Technology Update documents, asks an OpenAI model to
synthesize them into structured meeting minutes, and renders the result as
both a .docx and a .pdf.

The source documents are attached to the request as native OpenAI file inputs
(.pdf, .docx, .txt, .md, ... -- see NATIVE_MIME_TYPES).  PDFs are the best
input format: OpenAI sends the model both the extracted text and an image of
each page, which preserves the layout of financial tables.  Subtitle files
(.srt / .vtt) are not a file type OpenAI accepts, so those are flattened to
timestamped plain text locally and sent inline.

The model is asked for the minutes as structured JSON (see MINUTES_SCHEMA) and
this script renders the .docx and .pdf locally from that JSON, following the
layout of "FAC Meeting Minutes Template.pdf".

OpenAI's code interpreter tool could write the two documents itself, but
rendering locally is deliberate: the formatting is then deterministic and
identical from month to month, a layout fix is a code change rather than a
prompt change, and --from-json re-renders both documents without paying for
another inference.  The model supplies content; this script supplies form.
"""

import argparse
import base64
import json
import os
import re
import sys

from datetime import datetime

from openai import OpenAI

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                TableStyle)

#############################################################################
#
# Prompt.  This is an edited/tightened version of the original prompt,
# adapted to ask for structured JSON instead of asking the model to
# produce files itself.
#
#############################################################################

SYSTEM_PROMPT = """\
You are an expert administrative assistant who writes the minutes for the
Finance Advisory Committee ("FAC") of Epiphany Catholic Church ("ECC") in
Louisville, KY.

You will be given four source documents:

* Meeting Transcription: the raw text or transcript of the live discussion.
  It is typically machine-generated, so expect imperfect wording and
  punctuation.  Lines may be prefixed with a timestamp of the form
  [HH:MM:SS], measured from the START of the recording rather than from the
  wall clock, and with an anonymized speaker tag such as [SPEAKER_04].  Those
  tags are produced by automatic speaker diarization: they are consistent
  within the transcript but carry no names, so infer who each speaker is from
  context (self-introductions, people addressing each other by name, roles)
  and never present a raw SPEAKER_NN tag in the minutes.  Machine
  transcription also mangles proper names -- map them back to the committee
  members listed below when the intent is clear.
* Financial Highlights Document: a detailed summary of monthly income, budget
  variances, and expenditures.
* Financial Observations Document: a summary of noted observations regarding
  monthly income, budget variances, and expenditure highlights.
* Technology Update Document: a report detailing licensing, hardware
  refreshes, and infrastructure projects.  It may also be titled "ECC Tech
  Committee" or "ECC Tech Updates".

Synthesize them into a single set of professional meeting minutes, returned as
JSON matching the provided schema.  You produce content only; a separate
program renders the JSON into a Word document and a PDF, so do not concern
yourself with fonts, page layout, or file formats.

# DOCUMENT STRUCTURE

Emit the sections below, in this order, as entries in the "sections" array.
Omit a section entirely only if the source documents contain nothing at all
for it.

1. "Approval of Last Meeting's Minutes"
   The status of the review and approval of the prior month's minutes.  Note
   any corrections or changes made prior to approval, or state that the
   minutes were approved as submitted.

2. "Financial Report"
   A short lead-in paragraph naming the month under review, then a table with
   the columns: Category, Actual, Budget, YTD Actual.  Populate it from the
   Financial Highlights document (typical rows: General Collections, Hispanic
   Collections, Other Income, Total Income, Total Expenses).  Follow the table
   with a subheading "Financial Highlights & Observations" and a bulleted list
   drawn ONLY from the Financial Observations document, covering trends such
   as collections performance and any large individual expenses.

3. "Future Fiscal Planning & Budgeting"
   Highlights of any discussion of the next fiscal year's budget: revenue
   targets, staffing and cost-of-living adjustments (COLAs), and planned
   capital expenditures (CapEx).  Capture every relevant discussion point from
   the transcript.

4. "Facilities / Grounds Update"
   Physical plant projects, repairs, and efficiency measures, from the
   transcript or any facilities report.

5. "Technology Report"
   See the Technology Update rules below.

6. "Open Discussion"
   A summary of incidental topics from the transcript.  Be specific about the
   key discussion points.

7. "FAC Bulletin Highlights"
   See the Section 7 extraction rule below.

# CRITICAL PROCESSING RULES

* Source isolation: every item from the Technology Update / ECC Tech Committee
  document belongs EXCLUSIVELY in the Technology Report section.  Do not
  repeat tech infrastructure costs or software initiatives (migrations,
  hardware cycles) in the budget or facilities sections unless they are
  discussed there as a financial line item in the Financial Observations.

* No hallucination: include ONLY items found in the documents provided for
  this session.  Never carry over data (pricing, vendors, names, projects)
  from any other meeting or from your own general knowledge.

* Data integrity: preserve all specific dollar amounts and percentage
  variances exactly as they appear in the source documents.

* Tone: formal, administrative, and objective.  Past tense, third person.

# STRICT EXECUTION CONSTRAINTS

These rules are absolute; failing any of them is a failed response.

1. Meeting date
   Determine the date of the meeting from the source documents and report it
   in "meeting_date" as YYYY-MM-DD.  If it genuinely cannot be determined,
   return the empty string rather than guessing.

2. Logistics
   Report the meeting's start time and adjournment time (e.g. "7:01 PM").  If
   the adjournment time is not stated explicitly, calculate it by adding the
   transcript timestamp at which the meeting is adjourned (or, failing that,
   the final timestamp in the transcript) to the start time -- remember that
   transcript timestamps are elapsed time from the start of the recording, not
   clock times.  If a time cannot be determined at all, return the literal
   placeholder "[START TIME]" or "[ADJOURNMENT TIME]" so that a human can fill
   it in; never invent a time.

3. Attendance
   The committee's typical members are: Father Toan Do, Dave Gerwig, Doug
   Wolz, Sissy Watson, Dick Bowles, Mary Schumer, Keith Fowler, Suzanne Chase,
   Jay Fields, and Bill Stader.  Others may attend as guests.
   * Sort BOTH the "attendees" and "absentees" lists alphabetically by LAST
     name (Bowles, Chase, Fields, Fowler, Gerwig, Schumer, Stader, Toan Do,
     Watson, Wolz), and write each entry as the person's full name.
   * Cross-check both lists against the typical members above: every typical
     member should appear in exactly one of the two lists.
   * If attendance cannot be determined from the transcript, return the single
     placeholder entry "[ATTENDEES]" or "[ABSENTEES]" in the relevant list.

4. Technology Report comprehensive inclusion
   Include EVERY project, update, and initiative mentioned in the Technology
   Update document.  Do not summarize, merge, or drop items: a hardware order,
   a software license migration, a driver glitch, and an AI training plan each
   get their own bullet.  This rule overrides any general instinct toward
   brevity.  Give each bullet a short bold-style lead-in label followed by a
   colon (e.g. "UPS Failure: The uninterruptible power supply ...").

5. FAC Bulletin Highlights extraction
   This section is a concise 2-3 item list of the major takeaways the
   Committee agreed to report to the parish staff.
   * Extract it from the specific discussion near the END of the transcript
     where the Committee asks what should be reported to staff, or what goes
     in the bulletin highlights.
   * Ignore general financial observations and individual members' personal
     insights when building this list.
   * Do not substitute financial figures (such as collection totals) unless
     the transcript explicitly designates those figures as the bulletin
     highlights.
"""

#############################################################################
#
# JSON schema for the model's structured output.  OpenAI strict mode requires
# every property to be listed in "required" and additionalProperties: false,
# so "unused" fields are returned empty rather than omitted.
#
#############################################################################

BULLET_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "text": {
            "type": "string",
            "description": "Bullet text.  May begin with a short lead-in "
                           "label followed by a colon.",
        },
        "subbullets": {
            "type": "array",
            "description": "Nested bullets, or an empty array.",
            "items": {"type": "string"},
        },
    },
    "required": ["text", "subbullets"],
}

BLOCK_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "type": {
            "type": "string",
            "enum": ["paragraph", "subheading", "bullets", "table"],
        },
        "text": {
            "type": "string",
            "description": "Content for a paragraph or subheading block.  "
                           "Empty string for other block types.",
        },
        "bullets": {
            "type": "array",
            "description": "Content for a bullets block.  Empty array for "
                           "other block types.",
            "items": BULLET_SCHEMA,
        },
        "table": {
            "description": "Content for a table block.  Null for other block "
                           "types.",
            "anyOf": [
                {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "columns": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                        "rows": {
                            "type": "array",
                            "description": "Each row has one cell per column.",
                            "items": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                        },
                    },
                    "required": ["columns", "rows"],
                },
                {"type": "null"},
            ],
        },
    },
    "required": ["type", "text", "bullets", "table"],
}

SECTION_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "heading": {"type": "string"},
        "blocks": {
            "type": "array",
            "description": "The section's content, in the order it should "
                           "appear.",
            "items": BLOCK_SCHEMA,
        },
    },
    "required": ["heading", "blocks"],
}

MINUTES_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "meeting_date": {
            "type": "string",
            "description": "Date of the meeting as YYYY-MM-DD, or the empty "
                           "string if it cannot be determined.",
        },
        "start_time": {"type": "string"},
        "adjournment_time": {"type": "string"},
        "attendees": {
            "type": "array",
            "description": "Full names, sorted alphabetically by last name.",
            "items": {"type": "string"},
        },
        "absentees": {
            "type": "array",
            "description": "Full names, sorted alphabetically by last name.",
            "items": {"type": "string"},
        },
        "sections": {
            "type": "array",
            "items": SECTION_SCHEMA,
        },
    },
    "required": ["meeting_date", "start_time", "adjournment_time",
                 "attendees", "absentees", "sections"],
}

DOC_TITLE = "Finance Advisory Committee Meeting Minutes"

#############################################################################
#
# Input documents
#
# Most document types are attached to the request as native OpenAI file
# inputs; OpenAI does the text extraction (and, for PDFs, also sends the model
# an image of each page).  Subtitle files are not an accepted file type, so
# they are flattened locally and sent as inline text.
#
#############################################################################

NATIVE_MIME_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".odt": "application/vnd.oasis.opendocument.text",
    ".rtf": "application/rtf",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".csv": "text/csv",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

# OpenAI caps file inputs at 50 MB per file and 50 MB per request.
MAX_FILE_BYTES = 50 * 1024 * 1024

# Subtitle formats: flattened locally rather than attached.
FLATTEN_EXTENSIONS = (".srt", ".vtt")

SUPPORTED_EXTENSIONS = sorted(list(NATIVE_MIME_TYPES) + list(FLATTEN_EXTENSIONS))


def flatten_subtitles(filename):
    """Flatten an .srt/.vtt subtitle file into "[HH:MM:SS] text" lines.

    The cue numbers and the "-->" end times are noise for the model, but the
    start timestamps are load-bearing: the prompt uses them to work out the
    adjournment time.
    """

    time_re = re.compile(r"^(?:(\d{2}):)?(\d{2}):(\d{2})[,.]\d+\s*-->")

    out = []
    timestamp = None
    text = []

    def flush():
        if text:
            out.append("[%s] %s" % (timestamp or "??:??:??", " ".join(text)))
        del text[:]

    with open(filename, "r", encoding="utf-8-sig", errors="replace") as fp:
        for line in fp:
            line = line.strip()
            match = time_re.match(line)
            if match:
                flush()
                timestamp = "%s:%s:%s" % (match.group(1) or "00",
                                          match.group(2), match.group(3))
            elif not line:
                flush()
            elif line.isdigit() and not text:
                # Subtitle cue number; skip it.
                pass
            elif line.startswith("WEBVTT") or line.startswith("NOTE "):
                pass
            else:
                text.append(line)
    flush()

    return "\n".join(out)


def load_document(label, filename, log):
    """Load one source document into a content part for the Responses API.

    Returns a dict shaped for the API: either an "input_file" part carrying
    the file itself, or an "input_text" part carrying flattened text.
    """

    if not os.path.isfile(filename):
        raise ValueError('%s: no such file: "%s"' % (label, filename))

    ext = os.path.splitext(filename)[1].lower()
    size = os.path.getsize(filename)

    if ext in FLATTEN_EXTENSIONS:
        text = flatten_subtitles(filename).strip()
        if not text:
            raise ValueError('No text could be extracted from "%s"' % filename)
        log("Read %s (%s, flattened to %d characters of text)"
            % (filename, ext, len(text)))
        return {
            "type": "input_text",
            "text": "BEGIN %s (source file: %s)\n\n%s\n\nEND %s"
                    % (label, os.path.basename(filename), text, label),
        }

    mime = NATIVE_MIME_TYPES.get(ext)
    if mime is None:
        raise ValueError(
            'Do not know how to read "%s": unsupported file type "%s" '
            "(supported: %s)"
            % (filename, ext, ", ".join(SUPPORTED_EXTENSIONS)))

    if size > MAX_FILE_BYTES:
        raise ValueError(
            '"%s" is %.1f MB; OpenAI accepts at most %d MB per file'
            % (filename, size / 1024.0 / 1024.0,
               MAX_FILE_BYTES // 1024 // 1024))

    with open(filename, "rb") as fp:
        encoded = base64.b64encode(fp.read()).decode("ascii")

    log("Attaching %s (%s, %.1f KB)" % (filename, ext, size / 1024.0))

    return {
        "type": "input_file",
        "filename": os.path.basename(filename),
        "file_data": "data:%s;base64,%s" % (mime, encoded),
    }

#############################################################################
#
# OpenAI inference
#
#############################################################################

def build_input_content(documents):
    """Assemble the four source documents into one user message.

    documents is an ordered list of (label, filename, part) tuples.  Each
    attached file is preceded by a text part naming it, so that the model can
    tell the four documents apart.
    """

    content = [{
        "type": "input_text",
        "text": "Here are the source documents for this meeting.",
    }]

    for label, filename, part in documents:
        if part["type"] == "input_file":
            content.append({
                "type": "input_text",
                "text": "The next attached file is the %s (%s)."
                        % (label, os.path.basename(filename)),
            })
        content.append(part)

    content.append({
        "type": "input_text",
        "text": "Generate the meeting minutes as JSON conforming to the "
                "schema, following every rule in your instructions.",
    })

    return content


def describe_content(content):
    """Human-readable rendering of the request, for --save-prompt."""

    out = [SYSTEM_PROMPT, ""]

    for part in content:
        if part["type"] == "input_text":
            out.append(part["text"])
        else:
            out.append("<<< attached file: %s (%d bytes of base64) >>>"
                       % (part["filename"], len(part["file_data"])))
        out.append("")

    return "\n".join(out)


def generate_minutes(api_key, model, reasoning_effort, max_output_tokens,
                     content, log):
    """Send the request to OpenAI and return the minutes as a dict."""

    client = OpenAI(api_key=api_key)

    kwargs = {
        "model": model,
        "instructions": SYSTEM_PROMPT,
        "input": [{"role": "user", "content": content}],
        "max_output_tokens": max_output_tokens,
        "text": {
            "format": {
                "type": "json_schema",
                "name": "fac_meeting_minutes",
                "strict": True,
                "schema": MINUTES_SCHEMA,
            },
        },
    }

    # Reasoning effort only applies to the reasoning models (gpt-5, o*).
    if reasoning_effort and re.match(r"^(gpt-5|o\d)", model):
        kwargs["reasoning"] = {"effort": reasoning_effort}

    log("Sending request to OpenAI (model %s)..." % model)

    response = client.responses.create(**kwargs)

    if response.status == "incomplete":
        reason = getattr(response.incomplete_details, "reason", "unknown")
        raise RuntimeError(
            "OpenAI returned an incomplete response (%s).  If this is "
            "max_output_tokens, re-run with a larger --max-output-tokens."
            % reason)

    usage = response.usage
    if usage:
        log("OpenAI usage: %d input tokens, %d output tokens"
            % (usage.input_tokens, usage.output_tokens))

    return json.loads(response.output_text)

#############################################################################
#
# Shared rendering helpers
#
#############################################################################

def format_meeting_date(minutes, date_override):
    """Return (YYYY-MM-DD, "Month D, YYYY") for the meeting."""

    iso = date_override or minutes.get("meeting_date", "")

    try:
        when = datetime.strptime(iso, "%Y-%m-%d")
    except ValueError:
        raise ValueError(
            'Could not determine the meeting date (the model returned "%s").  '
            "Re-run with --date YYYY-MM-DD to set it explicitly." % iso)

    # No %-d on all platforms, so strip the leading zero by hand.
    pretty = "%s %d, %d" % (when.strftime("%B"), when.day, when.year)

    return iso, pretty


def format_names(names, placeholder):
    if not names:
        return placeholder
    return ", ".join(names)


def numeric_columns(table):
    """Indexes of the columns whose data cells look like numbers/currency."""

    out = set()

    for index in range(len(table["columns"])):
        cells = [row[index] for row in table["rows"] if index < len(row)]
        cells = [cell.strip() for cell in cells if cell.strip()]
        if cells and all(re.match(r"^[-+(]?[$]?[\d,.]+%?\)?$", cell)
                         for cell in cells):
            out.add(index)

    return out

#############################################################################
#
# .docx rendering
#
#############################################################################

DOCX_FONT = "Arial"


def _docx_shade_cell(cell, fill):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _docx_para(document, text, size=10.5, bold=False, align=None,
               space_after=6, style=None):
    para = document.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)

    if text:
        run = para.add_run(text)
        run.font.name = DOCX_FONT
        run.font.size = Pt(size)
        run.font.bold = bold

    return para


def _docx_bullets(document, bullets):
    for bullet in bullets:
        _docx_para(document, bullet["text"], style="List Bullet",
                   space_after=2)
        for sub in bullet.get("subbullets", []):
            _docx_para(document, sub, style="List Bullet 2", space_after=2)


def _docx_table(document, table):
    columns = table["columns"]
    rows = table["rows"]
    right = numeric_columns(table)

    doc_table = document.add_table(rows=1, cols=len(columns))
    doc_table.style = "Table Grid"
    doc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, name in enumerate(columns):
        cell = doc_table.rows[0].cells[index]
        cell.text = ""
        para = cell.paragraphs[0]
        if index in right:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(name)
        run.font.name = DOCX_FONT
        run.font.size = Pt(10)
        run.font.bold = True
        _docx_shade_cell(cell, "F1F1F1")

    for row in rows:
        cells = doc_table.add_row().cells
        for index in range(len(columns)):
            value = row[index] if index < len(row) else ""
            cell = cells[index]
            cell.text = ""
            para = cell.paragraphs[0]
            if index in right:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(value)
            run.font.name = DOCX_FONT
            run.font.size = Pt(10)

    document.add_paragraph()


def render_docx(minutes, pretty_date, filename):
    document = docx.Document()

    for section in document.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = DOCX_FONT
    normal.font.size = Pt(10.5)

    # Title block
    _docx_para(document, DOC_TITLE, size=13, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _docx_para(document, pretty_date, size=13, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # Logistics and attendance
    _docx_para(document, "Logistics: Start Time: %s | Adjournment Time: %s"
               % (minutes["start_time"], minutes["adjournment_time"]),
               space_after=12)
    _docx_para(document, "Attendance:  %s"
               % format_names(minutes["attendees"], "[ATTENDEES]"),
               space_after=6)
    _docx_para(document, "Absent:  %s"
               % format_names(minutes["absentees"], "[ABSENTEES]"),
               space_after=12)

    for section in minutes["sections"]:
        heading = _docx_para(document, section["heading"], size=17,
                             space_after=6)
        heading.paragraph_format.space_before = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        for block in section["blocks"]:
            kind = block["type"]
            if kind == "paragraph":
                _docx_para(document, block["text"])
            elif kind == "subheading":
                _docx_para(document, block["text"], size=12.5, space_after=4)
            elif kind == "bullets":
                _docx_bullets(document, block["bullets"])
            elif kind == "table" and block["table"]:
                _docx_table(document, block["table"])

    document.save(filename)

#############################################################################
#
# .pdf rendering
#
#############################################################################

PDF_FONT = "Helvetica"

PDF_STYLES = {
    "title": ParagraphStyle(
        "title", fontName=PDF_FONT + "-Bold", fontSize=13, leading=17,
        alignment=TA_CENTER),
    "heading": ParagraphStyle(
        "heading", fontName=PDF_FONT, fontSize=17, leading=21,
        spaceBefore=16, spaceAfter=8),
    "subheading": ParagraphStyle(
        "subheading", fontName=PDF_FONT, fontSize=12.5, leading=16,
        spaceBefore=6, spaceAfter=4),
    "body": ParagraphStyle(
        "body", fontName=PDF_FONT, fontSize=10.5, leading=14, spaceAfter=6),
    "bullet": ParagraphStyle(
        "bullet", fontName=PDF_FONT, fontSize=10.5, leading=14, spaceAfter=3,
        leftIndent=22, bulletIndent=8, bulletFontName=PDF_FONT,
        bulletFontSize=10.5),
    "subbullet": ParagraphStyle(
        "subbullet", fontName=PDF_FONT, fontSize=10.5, leading=14, spaceAfter=3,
        leftIndent=44, bulletIndent=30, bulletFontName=PDF_FONT,
        bulletFontSize=10.5),
    "cell": ParagraphStyle(
        "cell", fontName=PDF_FONT, fontSize=10, leading=13),
    "cell-right": ParagraphStyle(
        "cell-right", fontName=PDF_FONT, fontSize=10, leading=13,
        alignment=TA_RIGHT),
    "cell-head": ParagraphStyle(
        "cell-head", fontName=PDF_FONT + "-Bold", fontSize=10, leading=13),
    "cell-head-right": ParagraphStyle(
        "cell-head-right", fontName=PDF_FONT + "-Bold", fontSize=10,
        leading=13, alignment=TA_RIGHT),
}


def _pdf_escape(text):
    # ReportLab's Paragraph parses its input as mini-HTML.
    return (text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))


def _pdf_bullets(bullets):
    """Render a bullet list as individual bulleted paragraphs.

    ReportLab's ListFlowable is avoided here: it mutates the ParagraphStyle
    objects handed to it, which leaks indentation into every flowable that
    follows the list.
    """

    out = []

    for bullet in bullets:
        out.append(Paragraph(_pdf_escape(bullet["text"]),
                             PDF_STYLES["bullet"], bulletText="•"))
        for sub in bullet.get("subbullets", []):
            # "o", not a hollow-circle character: the base PDF fonts do not
            # have one, and ReportLab renders the missing glyph as a black box.
            out.append(Paragraph(_pdf_escape(sub), PDF_STYLES["subbullet"],
                                 bulletText="o"))

    return out


def _pdf_table(table, width):
    columns = table["columns"]
    rows = table["rows"]
    right = numeric_columns(table)

    data = [[Paragraph(_pdf_escape(name),
                       PDF_STYLES["cell-head-right" if index in right
                                  else "cell-head"])
             for index, name in enumerate(columns)]]

    for row in rows:
        data.append([
            Paragraph(_pdf_escape(row[index] if index < len(row) else ""),
                      PDF_STYLES["cell-right" if index in right else "cell"])
            for index in range(len(columns))
        ])

    # Give the first (label) column extra room; split the rest evenly.
    if len(columns) > 1:
        first = width * 0.34
        widths = [first] + [(width - first) / (len(columns) - 1)] * (len(columns) - 1)
    else:
        widths = [width]

    pdf_table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    pdf_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#999999")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F1F1F1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
    ]))

    return pdf_table


def render_pdf(minutes, pretty_date, filename):
    document = SimpleDocTemplate(
        filename, pagesize=letter,
        leftMargin=inch, rightMargin=inch,
        topMargin=inch, bottomMargin=inch,
        title="%s - %s" % (DOC_TITLE, pretty_date))

    width = document.width

    story = [
        Paragraph(_pdf_escape(DOC_TITLE), PDF_STYLES["title"]),
        Paragraph(_pdf_escape(pretty_date), PDF_STYLES["title"]),
        Spacer(1, 24),
        Paragraph(_pdf_escape("Logistics: Start Time: %s | Adjournment Time: %s"
                              % (minutes["start_time"],
                                 minutes["adjournment_time"])),
                  PDF_STYLES["body"]),
        Spacer(1, 6),
        Paragraph(_pdf_escape("Attendance:  %s"
                              % format_names(minutes["attendees"],
                                             "[ATTENDEES]")),
                  PDF_STYLES["body"]),
        Paragraph(_pdf_escape("Absent:  %s"
                              % format_names(minutes["absentees"],
                                             "[ABSENTEES]")),
                  PDF_STYLES["body"]),
    ]

    for section in minutes["sections"]:
        story.append(Paragraph(_pdf_escape(section["heading"]),
                               PDF_STYLES["heading"]))

        for block in section["blocks"]:
            kind = block["type"]
            if kind == "paragraph":
                story.append(Paragraph(_pdf_escape(block["text"]),
                                       PDF_STYLES["body"]))
            elif kind == "subheading":
                story.append(Paragraph(_pdf_escape(block["text"]),
                                       PDF_STYLES["subheading"]))
            elif kind == "bullets":
                story.extend(_pdf_bullets(block["bullets"]))
                story.append(Spacer(1, 4))
            elif kind == "table" and block["table"]:
                story.append(_pdf_table(block["table"], width))
                story.append(Spacer(1, 10))

    document.build(story)

#############################################################################
#
# Command line
#
#############################################################################

def read_api_key(filename):
    with open(filename, "r") as fp:
        key = fp.read().strip()

    if not key:
        raise ValueError('The OpenAI API key file "%s" is empty' % filename)

    return key


def setup_cli():
    parser = argparse.ArgumentParser(
        description="Generate FAC meeting minutes (.docx and .pdf) from a "
                    "meeting transcript and the monthly financial and "
                    "technology documents.")

    parser.add_argument("--openai-api-keyfile", required=True,
                        help="File containing the OpenAI API key")

    parser.add_argument("--transcript", required=True,
                        help="Meeting transcript (.txt, .md, .srt, .vtt, "
                             ".docx, or .pdf)")
    parser.add_argument("--financial-highlights", required=True,
                        help="Financial Highlights document")
    parser.add_argument("--financial-observations", required=True,
                        help="Financial Observations document")
    parser.add_argument("--technology-update", required=True,
                        help="Technology Update / ECC Tech Committee document")

    parser.add_argument("--model", default="gpt-5",
                        help="OpenAI model to use (default: %(default)s)")
    parser.add_argument("--reasoning-effort", default="medium",
                        choices=["minimal", "low", "medium", "high"],
                        help="Reasoning effort for reasoning models "
                             "(default: %(default)s)")
    parser.add_argument("--max-output-tokens", type=int, default=32000,
                        help="Cap on the model's output, including reasoning "
                             "tokens (default: %(default)s)")

    parser.add_argument("--outdir", default=".",
                        help="Directory for the output files "
                             "(default: the current directory)")
    parser.add_argument("--date",
                        help="Meeting date as YYYY-MM-DD, overriding the date "
                             "the model extracts from the documents")
    parser.add_argument("--basename",
                        help="Base filename for the outputs, overriding the "
                             'default of "YYYY-MM-DD FAC Meeting notes"')

    parser.add_argument("--save-json",
                        help="Also write the model's structured output to "
                             "this JSON file")
    parser.add_argument("--from-json",
                        help="Skip the OpenAI request and render from a "
                             "previously saved JSON file.  Useful when "
                             "iterating on formatting.")
    parser.add_argument("--save-prompt",
                        help="Write the assembled prompt to this file")
    parser.add_argument("--dry-run", action="store_true",
                        help="Read the inputs and assemble the prompt, but do "
                             "not call OpenAI or write any output documents")

    parser.add_argument("--quiet", action="store_true",
                        help="Suppress progress messages")

    args = parser.parse_args()

    if args.from_json and args.dry_run:
        parser.error("--from-json and --dry-run are mutually exclusive")

    return args


def main():
    args = setup_cli()

    def log(message):
        if not args.quiet:
            print(message)

    if args.from_json:
        with open(args.from_json, "r") as fp:
            minutes = json.load(fp)
        log("Rendering from %s (no OpenAI request)" % args.from_json)

    else:
        specs = [
            ("MEETING TRANSCRIPTION", args.transcript),
            ("FINANCIAL HIGHLIGHTS DOCUMENT", args.financial_highlights),
            ("FINANCIAL OBSERVATIONS DOCUMENT", args.financial_observations),
            ("TECHNOLOGY UPDATE DOCUMENT", args.technology_update),
        ]

        documents = [(label, filename, load_document(label, filename, log))
                     for label, filename in specs]

        attached = sum(os.path.getsize(filename)
                       for _, filename, part in documents
                       if part["type"] == "input_file")
        if attached > MAX_FILE_BYTES:
            raise ValueError(
                "The attached files total %.1f MB; OpenAI accepts at most "
                "%d MB per request"
                % (attached / 1024.0 / 1024.0,
                   MAX_FILE_BYTES // 1024 // 1024))

        content = build_input_content(documents)

        if args.save_prompt:
            with open(args.save_prompt, "w") as fp:
                fp.write(describe_content(content))
            log("Wrote prompt to %s" % args.save_prompt)

        if args.dry_run:
            log("Dry run: not calling OpenAI, not writing any documents")
            return

        minutes = generate_minutes(
            api_key=read_api_key(args.openai_api_keyfile),
            model=args.model,
            reasoning_effort=args.reasoning_effort,
            max_output_tokens=args.max_output_tokens,
            content=content,
            log=log)

        if args.save_json:
            with open(args.save_json, "w") as fp:
                json.dump(minutes, fp, indent=2)
            log("Wrote %s" % args.save_json)

    iso_date, pretty_date = format_meeting_date(minutes, args.date)

    basename = args.basename or "%s FAC Meeting notes" % iso_date

    if not os.path.isdir(args.outdir):
        os.makedirs(args.outdir)

    docx_filename = os.path.join(args.outdir, basename + ".docx")
    pdf_filename = os.path.join(args.outdir, basename + ".pdf")

    render_docx(minutes, pretty_date, docx_filename)
    log("Wrote %s" % docx_filename)

    render_pdf(minutes, pretty_date, pdf_filename)
    log("Wrote %s" % pdf_filename)


if __name__ == "__main__":
    try:
        main()
    except (ValueError, RuntimeError) as e:
        print("Error: %s" % e, file=sys.stderr)
        sys.exit(1)
